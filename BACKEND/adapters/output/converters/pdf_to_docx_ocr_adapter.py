import os
import platform
import shutil
import tempfile
from pathlib import Path

from domain.entities import ConversionError, FileFormat
from domain.ports.output.file_converter_port import FileConverterPort
from infrastructure.config import Settings


class PdfToDocxOcrAdapter(FileConverterPort):
    """OCR-based PDF → DOCX converter for scanned / image-based PDFs.

    **Workflow**
    1. Render each PDF page to a PNG image via PyMuPDF.
    2. Run Tesseract OCR on each image to extract text (with hOCR
       or plain text depending on availability).
    3. Assemble a ``.docx`` document with ``python-docx``.

    **Requirements**
    - ``pytesseract`` (Python) — ``pip install pytesseract``
    - ``python-docx`` (Python) — ``pip install python-docx``
    - `Tesseract OCR engine <https://github.com/tesseract-ocr/tesseract>`_
      installed on the system and available in ``PATH`` (or configured
      via ``FC_TESSERACT_PATH`` or ``settings.tesseract_path``).

    **Languages**
    By default Tesseract uses ``eng``. Pass ``language`` in options::

        {"language": "spa+eng"}
    """

    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._tesseract_cmd = self._find_tesseract()

    # ── Port interface ──────────────────────────────────────────────

    def supports(self, source: FileFormat, target: FileFormat) -> bool:
        return source == FileFormat.PDF and target == FileFormat.DOCX

    def convert(self, source_path: Path, target_path: Path, options: dict) -> None:
        self._ensure_dependencies()

        language = options.get("language", "eng")
        dpi = options.get("dpi", self._settings.default_image_dpi)
        page_range = options.get("pages")  # optional (start, end) 1-indexed

        try:
            doc = fitz.open(source_path)
        except Exception as e:
            raise ConversionError(f"Could not open PDF: {e}") from e

        if doc.page_count == 0:
            doc.close()
            raise ConversionError("The PDF contains no pages to OCR")

        page_indices = self._resolve_page_indices(page_range, doc.page_count)

        try:
            self._ocr_to_docx(doc, page_indices, target_path, language, dpi)
        finally:
            doc.close()

    # ── OCR pipeline ────────────────────────────────────────────────

    def _ocr_to_docx(
        self,
        doc: fitz.Document,
        page_indices: list[int],
        target_path: Path,
        language: str,
        dpi: int,
    ) -> None:
        # Import here so ImportError is caught at registration time
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        document = Document()

        for idx, page_num in enumerate(page_indices):
            page = doc[page_num]

            # ── Render page to image ─────────────────────────────
            pix = page.get_pixmap(dpi=dpi)
            img_bytes = pix.tobytes("png")

            # ── OCR via Tesseract ────────────────────────────────
            text = self._ocr_image(img_bytes, language)

            # ── Add page heading ─────────────────────────────────
            heading = document.add_heading(f"Page {page_num + 1}", level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ── Add extracted text ───────────────────────────────
            if text.strip():
                paragraph = document.add_paragraph(text)
                paragraph.style.font.size = Pt(11)
            else:
                document.add_paragraph("[No text detected on this page]")

            # Page break between pages (except last)
            if idx < len(page_indices) - 1:
                document.add_page_break()

        # ── Save ─────────────────────────────────────────────────
        document.save(str(target_path))

    def _ocr_image(self, image_data: bytes, language: str) -> str:
        """Run Tesseract on a PNG image and return extracted text."""
        try:
            import pytesseract

            if self._tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = self._tesseract_cmd

            # Write image to a temporary file (pytesseract works with
            # files more reliably than in-memory on all platforms).
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(image_data)
                tmp_path = tmp.name

            try:
                text: str = pytesseract.image_to_string(
                    tmp_path,
                    lang=language,
                    config="--psm 3",
                )
                return text
            finally:
                os.unlink(tmp_path)

        except ImportError as e:
            raise ConversionError(
                "pytesseract is not installed. Run: pip install pytesseract"
            ) from e

    # ── Helpers ──────────────────────────────────────────────────────

    def _ensure_dependencies(self) -> None:
        """Verify both pytesseract and python-docx are importable."""
        missing: list[str] = []
        try:
            # noqa: F841 — import check only
            import pytesseract  # noqa: F401
        except ImportError:
            missing.append("pytesseract")

        try:
            import docx  # noqa: F401
        except ImportError:
            missing.append("python-docx")

        if self._tesseract_cmd is None:
            missing.append("Tesseract engine (not found in PATH)")

        if not missing:
            return

        raise ConversionError(
            "OCR conversion is not available. Missing dependencies: "
            + ", ".join(missing)
            + (
                ". Install Python packages with: pip install pytesseract python-docx"
                ". Install Tesseract from https://github.com/tesseract-ocr/tesseract"
            )
        )

    def _find_tesseract(self) -> str | None:
        """Locate the Tesseract binary."""
        explicit = self._settings.tesseract_path
        if explicit and Path(explicit).exists():
            return explicit

        # Try common locations per platform
        system = platform.system()
        if system == "Windows":
            candidates = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ]
            for c in candidates:
                if Path(c).exists():
                    return c

        found = shutil.which("tesseract")
        if found:
            return found

        # Extra check on macOS via Homebrew
        if system == "Darwin":
            brew = "/opt/homebrew/bin/tesseract"
            if Path(brew).exists():
                return brew

        return None

    def _resolve_page_indices(
        self,
        page_range: tuple[int, int] | list[int] | None,
        total_pages: int,
    ) -> list[int]:
        if page_range is None:
            return list(range(total_pages))

        if isinstance(page_range, (tuple, list)) and len(page_range) == 2:
            start, end = page_range
            if start > end:
                start, end = end, start
            start = max(1, start) - 1
            end = min(total_pages, end)
            if start >= total_pages or end <= 0 or start >= end:
                raise ConversionError(
                    f"Page range {page_range[0]}-{page_range[1]} is outside "
                    f"the document (total pages: {total_pages})"
                )
            return list(range(start, end))

        if isinstance(page_range, list):
            indices = []
            for p in page_range:
                if p < 1 or p > total_pages:
                    raise ConversionError(
                        f"Page {p} is outside the document "
                        f"(total pages: {total_pages})"
                    )
                indices.append(p - 1)
            return indices

        return list(range(total_pages))
